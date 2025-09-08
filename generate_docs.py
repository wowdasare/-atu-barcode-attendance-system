#!/usr/bin/env python3
"""
ATU Barcode Attendance System Documentation Generator
Converts markdown documentation to professional DOCX format
"""

import os
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Installing required packages...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def create_cover_page(doc):
    """Create a professional cover page"""
    # University Logo placeholder
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run("🏛️ ACCRA TECHNICAL UNIVERSITY")
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue color
    title_run.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Main Title
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_title_run = main_title.add_run("ATU BARCODE STUDENT\nATTENDANCE SYSTEM")
    main_title_run.font.size = Pt(24)
    main_title_run.font.color.rgb = RGBColor(31, 41, 55)  # Dark gray
    main_title_run.bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Comprehensive Project Documentation")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)  # Medium gray
    subtitle_run.italic = True
    
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing
    
    # Version Info
    version_info = doc.add_paragraph()
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_info.add_run("Version 1.0")
    version_run.font.size = Pt(14)
    version_run.font.color.rgb = RGBColor(37, 99, 235)
    version_run.bold = True
    
    # Date
    date_info = doc.add_paragraph()
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_info.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Add several blank lines to center content vertically
    for _ in range(10):
        doc.add_paragraph()
    
    # Footer information
    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_info.add_run("Modern Web-Based QR Code Attendance Management System\nDesigned for Educational Institutions")
    footer_run.font.size = Pt(11)
    footer_run.font.color.rgb = RGBColor(107, 114, 128)
    footer_run.italic = True
    
    add_page_break(doc)

def setup_styles(doc):
    """Set up document styles"""
    styles = doc.styles
    
    # Heading 1 style
    heading1 = styles['Heading 1']
    heading1.font.size = Pt(18)
    heading1.font.color.rgb = RGBColor(37, 99, 235)
    heading1.font.bold = True
    
    # Heading 2 style  
    heading2 = styles['Heading 2']
    heading2.font.size = Pt(16)
    heading2.font.color.rgb = RGBColor(31, 41, 55)
    heading2.font.bold = True
    
    # Heading 3 style
    heading3 = styles['Heading 3']
    heading3.font.size = Pt(14)
    heading3.font.color.rgb = RGBColor(55, 65, 81)
    heading3.font.bold = True
    
    # Normal style
    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.name = 'Calibri'

def parse_markdown_to_docx(doc, markdown_file):
    """Parse markdown content and add to DOCX document"""
    with open(markdown_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    code_content = []
    
    for line in lines:
        line = line.rstrip()
        
        # Skip the main title (already in cover page)
        if line.startswith('# ATU Barcode Student Attendance System'):
            continue
        if line == '## Comprehensive Project Documentation':
            continue
        if line == '---' and len([l for l in lines[:10] if l.strip()]) < 5:
            continue
            
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_content:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_content))
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(9)
                    code_run.font.color.rgb = RGBColor(55, 65, 81)
                    code_para.style.name = 'Normal'
                code_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # Handle headings
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip() and line[0].isdigit() and '. ' in line:
            para = doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
        
        # Handle regular paragraphs
        elif line.strip():
            # Skip horizontal rules
            if line.strip() == '---':
                continue
                
            para = doc.add_paragraph()
            
            # Process inline formatting
            text = line
            
            # Handle text with colon-based formatting (e.g., "Feature Name: Description")
            if '::' in text:
                parts = text.split('::', 1)
                if len(parts) == 2:
                    # Bold the part before ::, normal the part after
                    title_run = para.add_run(parts[0] + ':')
                    title_run.bold = True
                    title_run.font.color.rgb = RGBColor(37, 99, 235)
                    para.add_run(' ' + parts[1])
                else:
                    para.add_run(text)
            elif '`' in text:
                parts = text.split('`')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        para.add_run(part)
                    else:
                        run = para.add_run(part)
                        run.font.name = 'Consolas'
                        run.font.color.rgb = RGBColor(37, 99, 235)
            else:
                para.add_run(text)
        else:
            # Empty line - add spacing
            doc.add_paragraph()

def add_table_of_contents(doc):
    """Add a table of contents page"""
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # TOC entries
    toc_entries = [
        ("1. Project Overview", "3"),
        ("2. System Architecture", "8"),
        ("3. Features and Functionality", "12"),
        ("4. Technology Stack", "18"),
        ("5. Installation Guide", "22"),
        ("6. User Manual", "28"),
        ("7. API Documentation", "35"),
        ("8. Database Schema", "40"),
        ("9. Security Features", "45"),
        ("10. Deployment Guide", "48"),
        ("11. Troubleshooting", "52"),
        ("12. Future Enhancements", "58"),
        ("13. Technical Support", "62"),
        ("Appendices", "64")
    ]
    
    for entry, page in toc_entries:
        toc_para = doc.add_paragraph()
        toc_para.add_run(entry)
        toc_para.add_run('\t' * 5)  # Tab spacing
        page_run = toc_para.add_run(page)
        page_run.bold = True
    
    add_page_break(doc)

def create_documentation():
    """Main function to create the documentation"""
    print("🔧 Generating ATU Barcode Attendance System Documentation...")
    
    # Create new document
    doc = Document()
    
    # Setup document styles
    setup_styles(doc)
    
    # Create cover page
    print("📄 Creating cover page...")
    create_cover_page(doc)
    
    # Add table of contents
    print("📋 Adding table of contents...")
    add_table_of_contents(doc)
    
    # Parse and add markdown content
    print("📝 Converting markdown content...")
    markdown_file = 'ATU_Barcode_Attendance_System_Documentation.md'
    
    if os.path.exists(markdown_file):
        parse_markdown_to_docx(doc, markdown_file)
    else:
        print(f"❌ Error: {markdown_file} not found!")
        return
    
    # Save document
    output_file = f'ATU_Barcode_Attendance_System_Documentation_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(output_file)
    
    print(f"✅ Documentation generated successfully!")
    print(f"📁 File saved as: {output_file}")
    print(f"📊 File size: {os.path.getsize(output_file) / 1024:.1f} KB")
    
    return output_file

if __name__ == "__main__":
    try:
        output_file = create_documentation()
        
        # Print summary
        print("\n" + "="*60)
        print("📚 ATU BARCODE ATTENDANCE SYSTEM DOCUMENTATION")
        print("="*60)
        print("✅ Status: Successfully Generated")
        print(f"📁 Location: {os.path.abspath(output_file)}")
        print(f"📅 Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("\n📋 Document Contents:")
        print("   • Project Overview & Architecture")
        print("   • Features & Technology Stack")
        print("   • Installation & Deployment Guide")
        print("   • User Manual & API Documentation")
        print("   • Database Schema & Security")
        print("   • Troubleshooting & Future Plans")
        print("="*60)
        
    except Exception as e:
        print(f"❌ Error generating documentation: {e}")
        print("💡 Make sure you have the required packages installed:")
        print("   pip install python-docx")